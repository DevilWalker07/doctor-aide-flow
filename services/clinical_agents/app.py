import os
import uuid
import json
import asyncio
import tempfile
import io
import traceback
import base64
from datetime import datetime, timedelta
from typing import Dict, Any, Optional
from dotenv import load_dotenv

# Load environment variables from .env file if it exists
load_dotenv()

import uvicorn
from fastapi import FastAPI, File, UploadFile, BackgroundTasks, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from PIL import Image, ImageOps
import pillow_heif
pillow_heif.register_heif_opener()
import mammoth
import pypdfium2 as pdfium
from openai import AsyncOpenAI
from supabase import create_client
try:
    from supabase import Client
except ImportError:
    from supabase.client import Client

# Conditional import for docling as it's a heavy dependency
try:
    from docling.document_converter import DocumentConverter
    HAS_DOCLING = True
except ImportError:
    HAS_DOCLING = False
    print("Warning: docling not installed. Falling back to other extraction methods.")

app = FastAPI(title="Doutor Ajuda - Async Clinical Agents")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Supabase Job Storage
SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY", "")

supabase_client: Optional[Client] = None
if SUPABASE_URL and SUPABASE_SERVICE_KEY:
    try:
        supabase_client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
    except Exception as e:
        print(f"Failed to initialize Supabase: {e}")

jobs_memory = {}  # fallback se Supabase não estiver configurado

# Use OPENAI_API_KEY or VITE_OPENAI_API_KEY
openai_key = os.environ.get("OPENAI_API_KEY") or os.environ.get("VITE_OPENAI_API_KEY")
client = AsyncOpenAI(api_key=openai_key)

# Add heif/heic support
pillow_heif.register_heif_opener()

# PDF page cap for image fallback — protects context window and timeout
MAX_PDF_PAGES = int(os.environ.get("MAX_PDF_PAGES", "15"))

def save_job(job_id: str, data: dict):
    if supabase_client:
        try:
            # Check if exists
            existing = supabase_client.table("extraction_jobs") \
                .select("id").eq("job_id", job_id).execute()
            if existing.data:
                supabase_client.table("extraction_jobs") \
                    .update(data).eq("job_id", job_id).execute()
            else:
                supabase_client.table("extraction_jobs") \
                    .insert({"job_id": job_id, **data}).execute()
            return
        except Exception as e:
            print(f"Supabase save_job error: {e}")
    jobs_memory[job_id] = data

def get_job(job_id: str) -> Optional[dict]:
    if supabase_client:
        try:
            result = supabase_client.table("extraction_jobs") \
                .select("*").eq("job_id", job_id).execute()
            if result.data:
                return result.data[0]
        except Exception as e:
            print(f"Supabase get_job error: {e}")
    return jobs_memory.get(job_id)

def update_job(job_id: str, status: str, stage: str,
               result: dict = None, error: str = None):
    data = {"status": status, "stage": stage}
    if result is not None:
        data["result"] = result
    if error is not None:
        data["error"] = error
    save_job(job_id, data)

async def process_document_background(job_id: str, file_bytes: bytes, filename: str, content_type: str, model_override: Optional[str] = None):
    ext = os.path.splitext(filename)[1].lower()
    try:
        update_job(job_id, "processing", "Identificando tipo de arquivo...")
        
        extracted_text = ""
        image_base64_list = []

        if ext in [".heic", ".heif", ".jpg", ".jpeg", ".png", ".webp"]:
            update_job(job_id, "processing", "Preparando imagem...")
            image = Image.open(io.BytesIO(file_bytes))
            # Rotate per EXIF so phone photos arrive upright at the OpenAI Vision call.
            image = ImageOps.exif_transpose(image)
            # Resize if too large
            max_size = 1800
            if max(image.size) > max_size:
                ratio = max_size / max(image.size)
                image = image.resize((int(image.size[0] * ratio), int(image.size[1] * ratio)), Image.LANCZOS)
            
            output = io.BytesIO()
            if image.mode in ("RGBA", "P"):
                image = image.convert("RGB")
            image.save(output, format="JPEG", quality=88)
            image_base64_list.append(base64.b64encode(output.getvalue()).decode("utf-8"))

        elif ext in [".txt", ".md"]:
            update_job(job_id, "processing", "Lendo texto...")
            extracted_text = file_bytes.decode("utf-8", errors="replace")

        elif ext == ".docx":
            update_job(job_id, "processing", "Lendo documento...")
            # Tentar docling primeiro
            try:
                if HAS_DOCLING:
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                        tmp.write(file_bytes)
                        tmp_path = tmp.name
                    
                    converter = DocumentConverter()
                    doc = converter.convert(tmp_path)
                    extracted_text = doc.document.export_to_markdown()
                    os.remove(tmp_path)
                else:
                    raise ImportError("docling not available")
            except Exception as e:
                # Fallback to mammoth se docling falhar
                print("Docling falhou para DOCX, usando mammoth:", e)
                result = mammoth.extract_raw_text(io.BytesIO(file_bytes))
                extracted_text = result.value
                
            if not extracted_text.strip():
                raise ValueError("Não foi possível extrair texto do DOCX.")

        elif ext == ".pdf":
            update_job(job_id, "processing", "Lendo PDF...")
            try:
                if HAS_DOCLING:
                    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                        tmp.write(file_bytes)
                        tmp_path = tmp.name

                    converter = DocumentConverter()
                    doc = converter.convert(tmp_path)
                    extracted_text = doc.document.export_to_markdown()
                    os.remove(tmp_path)
                else:
                    extracted_text = ""
            except Exception as e:
                print("Docling falhou para PDF:", e)
                extracted_text = ""

            # If docling produced no usable text, try pypdfium2 text extraction
            # before falling back to vision (which is slow + token-heavy).
            if len(extracted_text.strip()) < 100:
                update_job(job_id, "processing", "Extraindo texto do PDF...")
                try:
                    pdf = pdfium.PdfDocument(file_bytes)
                    text_parts = []
                    for i in range(len(pdf)):
                        textpage = pdf[i].get_textpage()
                        text_parts.append(textpage.get_text_range())
                    extracted_text = "\n\n".join(t for t in text_parts if t).strip()
                except Exception as e:
                    print("pypdfium2 text extraction falhou:", e)
                    extracted_text = ""

            # Still no text? Render pages to images for OpenAI Vision — but cap
            # at MAX_PDF_PAGES to protect context window and request timeout.
            if len(extracted_text.strip()) < 100:
                update_job(job_id, "processing", "Processando páginas (OCR)...")
                pdf = pdfium.PdfDocument(file_bytes)
                total_pages = len(pdf)
                pages_to_render = min(total_pages, MAX_PDF_PAGES)
                for i in range(pages_to_render):
                    page = pdf[i]
                    bitmap = page.render(scale=2.5) # Aumentando escala para melhor leitura
                    pil_image = bitmap.to_pil()
                    if pil_image.mode in ("RGBA", "P"):
                        pil_image = pil_image.convert("RGB")
                    
                    max_size = 1800
                    if max(pil_image.size) > max_size:
                        ratio = max_size / max(pil_image.size)
                        pil_image = pil_image.resize((int(pil_image.size[0] * ratio), int(pil_image.size[1] * ratio)), Image.LANCZOS)

                    output = io.BytesIO()
                    pil_image.save(output, format="JPEG", quality=88)
                    image_base64_list.append(base64.b64encode(output.getvalue()).decode("utf-8"))
                
                extracted_text = ""
                if total_pages > pages_to_render:
                    print(f"PDF truncado: enviando {pages_to_render} de {total_pages} páginas para a IA.")

        else:
            raise Exception(f"Formato não suportado: {ext}")

        update_job(job_id, "processing", "Organizando dados clínicos com IA...")

        # ETAPA 3 — JSON CLÍNICO PADRONIZADO alignment
        prompt = f"""
Você é um médico assistente sênior treinado para extrair, de forma EXAUSTIVA e
PRECISA, todos os dados clínicos relevantes de um documento médico (admissão,
evolução, prescrição, sumário, transferência, etc).

CONTEXTO:
- Nome do arquivo: {filename}
- Data de hoje: {datetime.now().strftime("%Y-%m-%d")}

DOCUMENTO COMPLETO:
\"\"\"
{extracted_text}
\"\"\"

PROTOCOLO DE EXTRAÇÃO (siga em ordem):
1. LEIA o documento INTEIRO antes de começar. Não extraia campo a campo isoladamente.
2. Para CADA seção do documento (identificação, HDA, antecedentes, medicações,
   exame físico, exames, conduta, pendências, evolução), identifique a qual chave
   do JSON ela pertence e extraia TUDO o que estiver presente.
3. Se uma informação aparece em mais de um lugar, consolide. Não duplique.
4. Se houver datas em formatos brasileiros ("15/05/2026"), converta para ISO ("2026-05-15").
5. Para listas (problemas, ATB, medicações, pendências) extraia CADA item separadamente,
   não junte vários numa mesma string.
6. NUNCA invente. Se não encontrou, use null para escalares ou [] para listas e
   adicione o campo em `campos_incertos`.

ATENÇÃO ESPECIAL — campos onde modelos costumam falhar:
- IDADE: procure por "X anos", "X a", data de nascimento (calcule), "idoso de X".
- SEXO: procure pronomes ("o paciente", "a paciente"), "sexo M/F", "masc/fem".
- PROBLEMAS ATIVOS vs RESOLVIDOS: ATIVOS é o problem list em curso; RESOLVIDOS
  são condições que já foram tratadas (geralmente em "antecedentes" ou em
  "problemas resolvidos"). Separe.
- ANTECEDENTES PATOLÓGICOS: comorbidades crônicas (HAS, DM, DRC, ICC, neoplasia).
  Vão em `antecedentes_patologicos`, NÃO em `problemas_ativos`.
- MEDICAÇÕES DE USO CONTÍNUO (domiciliares, prévias) → `medicacoes_uso_continuo`.
- PRESCRIÇÃO ATUAL DA INTERNAÇÃO → `prescricao_internacao`.
- ANTIBIÓTICOS: extraia COM data_inicio, dose, via, frequência. Se não há data,
  use null e marque em campos_incertos.
  NORMALIZE o nome para esta lista canônica quando reconhecer o ATB:
  CEFTRIAXONA, PIPERACILINA + TAZOBACTAM, VANCOMICINA, MEROPENEM,
  AMOXICILINA + CLAVULANATO, AZITROMICINA, CLINDAMICINA, METRONIDAZOL,
  CIPROFLOXACINO, LEVOFLOXACINO, CEFEPIME, POLIMIXINA B, OXACILINA, GENTAMICINA.
  Apelidos comuns (PIPTAZO, PIP-TAZO, TAZOCIN, ROCEFIN, VANCO, MERO, CLAVULIN,
  CIPRO, LEVO, CLINDA, METRO, FLAGYL, GENTA) devem ser convertidos para a
  forma canônica acima.
- EXAMES DE IMAGEM (TC, RM, RX, USG, ECO) vão em `exames_imagem`, NÃO em laboratorios.
- LABORATÓRIOS: extraia em texto compacto inline (ex: "HB 11 | LEUCO 9.500 | PCR 35").
- EXAME FÍSICO: preencha cada sistema separadamente (geral/acv/ar/abdome/neuro/extremidades/pele).

ESQUEMA JSON OBRIGATÓRIO (retorne APENAS este objeto, sem markdown, sem comentários):

{{
  "nome": "string ou null",
  "idade": "number ou null",
  "sexo": "M | F | NI",
  "leito": "string ou null",
  "setor": "string ou null",
  "data_admissao": "YYYY-MM-DD ou null",
  "data_documento": "YYYY-MM-DD ou null",
  "motivo_admissao": "string ou null",
  "hda": "string completa ou null",
  "antecedentes_patologicos": ["string", "..."],
  "medicacoes_uso_continuo": ["string", "..."],
  "alergias": ["string", "..."],
  "problemas_ativos": ["string", "..."],
  "problemas_resolvidos": ["string", "..."],
  "lista_de_problemas": ["string", "..."],
  "prescricao_internacao": ["string", "..."],
  "medicacoes": ["string", "..."],
  "antibioticos": [
    {{ "nome": "string", "dose": "string", "via": "string", "frequencia": "string", "data_inicio": "YYYY-MM-DD ou null" }}
  ],
  "laboratorios": [
    {{ "data": "YYYY-MM-DD", "texto_compacto": "HB 11 | LEUCO 9.500 | ...", "valores": {{}} }}
  ],
  "exames_imagem": [
    {{ "data": "YYYY-MM-DD ou null", "modalidade": "TC/RM/RX/USG/ECO", "descricao": "string", "laudo": "string ou null" }}
  ],
  "exame_fisico_detalhado": {{
    "geral": "string ou null",
    "acv": "string ou null",
    "ar": "string ou null",
    "abdome": "string ou null",
    "neuro": "string ou null",
    "extremidades": "string ou null",
    "pele": "string ou null"
  }},
  "condutas": ["string", "..."],
  "pendencias": ["string", "..."],
  "alertas": ["string", "..."],
  "campos_incertos": ["string", "..."]
}}

REGRAS FINAIS:
- `lista_de_problemas` deve conter problemas_ativos + problemas_resolvidos (compatibilidade).
- Se o nome não estiver no texto, infira do arquivo (ex: "L01-JOAO-SILVA.pdf" → "JOAO SILVA").
- Cite SEMPRE em `campos_incertos` aqueles cuja extração foi por inferência ou está duvidosa.
- Antes de finalizar, REVISE mentalmente: "deixei algum dado de fora?". Se sim, inclua.
"""

        messages = [
            {"role": "system", "content": "Você é um médico especialista em transcrição de dados estruturados. Retorne apenas JSON válido e preciso."}
        ]

        if image_base64_list:
            content = [{"type": "text", "text": prompt}]
            for b64 in image_base64_list:
                content.append({"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}})
            messages.append({"role": "user", "content": content})
        else:
            messages.append({
                "role": "user",
                "content": f"{prompt}\n\nDocumento:\n{extracted_text}"
            })

        # Extração clínica precisa do modelo completo (gpt-4o), não da versão mini.
        # gpt-4o-mini perde nuances clínicas e omite campos em documentos longos.
        # Se OPENAI_MODEL apontar para um modelo inválido (ex: "gpt-4.1-mini"),
        # cai para gpt-4o (não para gpt-4o-mini).
        # Default em gpt-4o-mini (mais barato). Para ativar gpt-4o full,
        # basta setar OPENAI_MODEL=gpt-4o nas env vars do Railway.
        VALID_MODELS = {"gpt-4o", "gpt-4o-2024-11-20", "gpt-4o-2024-08-06", "gpt-4-turbo", "gpt-4o-mini"}
        model_name = model_override or os.environ.get("OPENAI_MODEL", "gpt-4o-mini")
        if model_name not in VALID_MODELS:
            print(f"[extracao] modelo='{model_name}' inválido, usando gpt-4o-mini.")
            model_name = "gpt-4o-mini"

        response = None
        for tentativa in range(3):
            try:
                response = await client.chat.completions.create(
                    model=model_name,
                    messages=messages,
                    temperature=0.0,
                    max_tokens=16000,
                    response_format={ "type": "json_object" }
                )
                break
            except Exception as e:
                if tentativa == 2:
                    raise e
                print(f"Erro na OpenAI (tentativa {tentativa+1}/3): {e}")
                await asyncio.sleep(2)

        result_content = response.choices[0].message.content
        
        try:
            clinical_json = json.loads(result_content)
        except Exception:
            clinical_json = {"error": "Falha ao parsear JSON", "raw": result_content}

        clinical_json["fileName"] = filename
        clinical_json["engine"] = "openai-vision" if image_base64_list else "openai-text"
        
        update_job(job_id, "done", "Pronto para revisão", result=clinical_json)

    except Exception as e:
        print(traceback.format_exc())
        update_job(job_id, "error", "Erro na extração", error=str(e))

# --- BULK EXTRACTION (passagem de plantão a partir de múltiplos documentos) ---

MAX_BULK_FILES = 20

bulk_jobs_memory: Dict[str, dict] = {}


def _save_bulk_job(job_id: str, data: dict):
    bulk_jobs_memory[job_id] = data


def _get_bulk_job(job_id: str) -> Optional[dict]:
    return bulk_jobs_memory.get(job_id)


def _update_bulk_job(job_id: str, **fields):
    job = bulk_jobs_memory.get(job_id)
    if not job:
        return
    job.update(fields)
    job["updated_at"] = datetime.now().isoformat()


async def _generate_passagem_consolidada(patients: List[dict]) -> dict:
    """
    Recebe a lista de JSONs clínicos já extraídos (por gpt-4o-mini) e usa o
    gpt-4o full numa única chamada para ranquear gravidade, sugerir altas e
    montar a passagem de plantão consolidada.
    """
    if not patients:
        return {"resumo": "", "ranking": [], "alertas_globais": []}

    # Resumo mínimo por paciente para caber no contexto sem mandar TUDO.
    compact = []
    for p in patients:
        compact.append({
            "fileName": p.get("fileName"),
            "nome": p.get("nome"),
            "idade": p.get("idade"),
            "sexo": p.get("sexo"),
            "leito": p.get("leito"),
            "setor": p.get("setor"),
            "data_admissao": p.get("data_admissao"),
            "motivo_admissao": p.get("motivo_admissao"),
            "problemas_ativos": p.get("problemas_ativos") or p.get("lista_de_problemas") or [],
            "antibioticos": p.get("antibioticos") or [],
            "laboratorios": (p.get("laboratorios") or [])[:2],
            "exame_fisico": p.get("exame_fisico_detalhado") or {},
            "pendencias": p.get("pendencias") or [],
            "alertas": p.get("alertas") or [],
        })

    prompt = f"""
Você é um médico sênior responsável por preparar a PASSAGEM DE PLANTÃO para
o colega que está chegando. Recebeu abaixo os dados de {len(patients)} pacientes
já extraídos. Sua tarefa:

1) RANQUEAR todos por GRAVIDADE clínica:
   - "critico": VM, DVA, choque, sepse grave, instabilidade hemodinâmica
   - "grave": ATB IV ativo recente, dispneia, instabilidade, exame físico alterado
   - "moderado": estável mas em investigação ativa
   - "leve": estável, possível alta

2) Para cada paciente, indicar:
   - dias_de_internacao (a partir de data_admissao)
   - dias_de_antibiotico (D{{N}} a partir do ATB mais antigo ativo)
   - resumo_uma_linha (motivo + estado atual)
   - criterios_gravidade_atendidos: lista do que justifica a classificação
   - sugere_alta (boolean) e por_que_alta (se aplicável)
   - alertas_pendencias: pendências críticas a observar

3) Escrever uma passagem_texto consolidada e legível em PT-BR, em CAIXA ALTA,
   listando os pacientes da MAIOR PRA MENOR gravidade, com este formato por paciente:
       LEITO X — NOME (IDADEa SEXO) — D{{INTERNAÇÃO}}
       MOTIVO: ...
       ATB: ... D{{DIA}}
       LAB: ...
       PENDÊNCIAS: ...
       ALERTAS: ...

DADOS DOS PACIENTES:
{json.dumps(compact, ensure_ascii=False, indent=2)}

DATA DE HOJE: {datetime.now().strftime("%Y-%m-%d")}

REGRAS:
- Não invente dados. Se faltar, escreva "NÃO REFERIDO".
- Use SOMENTE os dados fornecidos acima.
- Antes de finalizar, garanta que TODOS os {len(patients)} pacientes aparecem no ranking.

RETORNE APENAS UM JSON com este schema:
{{
  "ranking": [
    {{
      "fileName": "string",
      "leito": "string",
      "nome": "string",
      "gravidade": "critico | grave | moderado | leve",
      "dias_de_internacao": "number ou null",
      "dias_de_antibiotico": "number ou null",
      "resumo_uma_linha": "string",
      "criterios_gravidade_atendidos": ["..."],
      "sugere_alta": false,
      "por_que_alta": "string ou null",
      "alertas_pendencias": ["..."]
    }}
  ],
  "passagem_texto": "string completa em caixa alta",
  "alertas_globais": ["alertas para o plantão como um todo"],
  "contagem": {{"critico": 0, "grave": 0, "moderado": 0, "leve": 0}}
}}
"""

    # Default em gpt-4o-mini para evitar custo extra enquanto OPENAI_ANALYSIS_MODEL
    # não estiver configurado. Para usar gpt-4o full na análise consolidada,
    # setar OPENAI_ANALYSIS_MODEL=gpt-4o no Railway.
    VALID_ANALYSIS_MODELS = {"gpt-4o", "gpt-4o-2024-11-20", "gpt-4o-2024-08-06", "gpt-4o-mini"}
    analysis_model = os.environ.get("OPENAI_ANALYSIS_MODEL", "gpt-4o-mini")
    if analysis_model not in VALID_ANALYSIS_MODELS:
        analysis_model = "gpt-4o-mini"

    response = await client.chat.completions.create(
        model=analysis_model,
        messages=[
            {"role": "system", "content": "Você é um médico sênior. Retorne apenas JSON válido."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.0,
        max_tokens=8000,
        response_format={"type": "json_object"},
    )
    raw = response.choices[0].message.content
    try:
        return json.loads(raw)
    except Exception:
        return {"error": "Falha ao parsear JSON da análise", "raw": raw}


async def process_bulk_background(bulk_job_id: str, files_data: List[tuple]):
    """
    files_data: lista de tuplas (file_bytes, filename, content_type).
    Cria N sub-jobs reusando process_document_background e roda em paralelo,
    depois consolida a passagem com gpt-4o.
    """
    try:
        sub_job_ids: List[str] = []
        for _, filename, _ct in files_data:
            sj = str(uuid.uuid4())
            save_job(sj, {
                "status": "queued",
                "stage": "Aguardando",
                "file_name": filename,
                "result": None,
                "error": None,
            })
            sub_job_ids.append(sj)

        _update_bulk_job(
            bulk_job_id,
            status="processing",
            stage=f"Extraindo {len(sub_job_ids)} documentos em paralelo...",
            sub_jobs=sub_job_ids,
            total=len(sub_job_ids),
        )

        # Roda todas as extrações em paralelo, FORÇANDO gpt-4o-mini por arquivo
        # para custo baixo. A análise consolidada usa gpt-4o full (1 chamada).
        bulk_extract_model = os.environ.get("OPENAI_BULK_MODEL", "gpt-4o-mini")
        tasks = [
            process_document_background(sj, b, fn, ct, model_override=bulk_extract_model)
            for sj, (b, fn, ct) in zip(sub_job_ids, files_data)
        ]
        await asyncio.gather(*tasks, return_exceptions=True)

        patients: List[dict] = []
        errors: List[dict] = []
        for sj in sub_job_ids:
            j = get_job(sj)
            if j and j.get("status") == "done" and j.get("result"):
                patients.append(j["result"])
            else:
                errors.append({
                    "sub_job_id": sj,
                    "file_name": (j or {}).get("file_name"),
                    "error": (j or {}).get("error") or "Falha desconhecida",
                })

        _update_bulk_job(
            bulk_job_id,
            stage=f"Gerando passagem consolidada de {len(patients)} pacientes...",
            extracted_count=len(patients),
            error_count=len(errors),
        )

        passagem = {}
        if patients:
            try:
                passagem = await _generate_passagem_consolidada(patients)
            except Exception as e:
                passagem = {"error": f"Falha na análise consolidada: {e}"}

        _update_bulk_job(
            bulk_job_id,
            status="done",
            stage="Concluído",
            patients=patients,
            passagem=passagem,
            errors=errors,
        )
    except Exception as e:
        print(traceback.format_exc())
        _update_bulk_job(bulk_job_id, status="error", error=str(e))


@app.post("/extract-bulk")
@app.post("/api/extract/bulk")
async def extract_bulk(background_tasks: BackgroundTasks, files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(status_code=400, detail="Nenhum arquivo enviado.")
    if len(files) > MAX_BULK_FILES:
        raise HTTPException(status_code=400, detail=f"Máximo de {MAX_BULK_FILES} arquivos por lote.")

    files_data = []
    for f in files:
        files_data.append((await f.read(), f.filename, f.content_type))

    bulk_job_id = str(uuid.uuid4())
    _save_bulk_job(bulk_job_id, {
        "status": "queued",
        "stage": "Lote recebido",
        "total": len(files_data),
        "file_names": [fn for _, fn, _ in files_data],
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat(),
        "sub_jobs": [],
        "patients": [],
        "passagem": None,
        "errors": [],
    })

    background_tasks.add_task(process_bulk_background, bulk_job_id, files_data)
    return {"bulk_job_id": bulk_job_id, "status": "queued", "total": len(files_data)}


@app.get("/extract-bulk/job/{bulk_job_id}")
@app.get("/api/extract/bulk/job/{bulk_job_id}")
async def get_bulk_job_status(bulk_job_id: str):
    job = _get_bulk_job(bulk_job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Lote não encontrado")

    # Inclui progresso real consultando sub-jobs.
    sub_ids = job.get("sub_jobs") or []
    progress = {"done": 0, "error": 0, "processing": 0, "queued": 0}
    sub_status = []
    for sj in sub_ids:
        s = get_job(sj)
        if not s:
            continue
        st = s.get("status", "queued")
        progress[st] = progress.get(st, 0) + 1
        sub_status.append({
            "sub_job_id": sj,
            "file_name": s.get("file_name"),
            "status": st,
            "stage": s.get("stage"),
        })

    return {**job, "progress": progress, "sub_status": sub_status}


# --- DOCUMENT EXPORT (DOCX / PDF / TXT) ---

class ExportRequest(BaseModel):
    text: str
    format: str  # "docx" | "pdf" | "txt"
    header: Optional[Dict[str, Any]] = None  # {hospital, medico, crm, paciente, leito, data, tipo}
    filename_hint: Optional[str] = None


def _safe_filename(s: str) -> str:
    if not s:
        return "evolucao"
    keep = "".join(c if c.isalnum() or c in "-_" else "_" for c in s)
    return keep.strip("_") or "evolucao"


def _build_docx_bytes(text: str, header: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    hospital = header.get("hospital") or ""
    medico = header.get("medico") or ""
    crm = header.get("crm") or ""
    paciente = header.get("paciente") or ""
    leito = header.get("leito") or ""
    data = header.get("data") or datetime.now().strftime("%d/%m/%Y")
    tipo = header.get("tipo") or "EVOLUÇÃO MÉDICA"

    if hospital or medico:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_left = h.add_run(hospital)
        run_left.bold = True
        h.add_run("\t" * 6)
        right = f"{medico}" + (f" — CRM {crm}" if crm else "")
        h.add_run(right)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(tipo.upper())
    title_run.bold = True
    title_run.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta_text = []
    if paciente: meta_text.append(f"PACIENTE: {paciente}")
    if leito: meta_text.append(f"LEITO: {leito}")
    if data: meta_text.append(f"DATA: {data}")
    meta.add_run(" · ".join(meta_text)).bold = True

    doc.add_paragraph("─" * 80)

    for raw_line in text.split("\n"):
        p = doc.add_paragraph(raw_line)
        for run in p.runs:
            run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf_bytes(text: str, header: Dict[str, Any]) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.colors import HexColor

    hospital = header.get("hospital") or ""
    medico = header.get("medico") or ""
    crm = header.get("crm") or ""
    paciente = header.get("paciente") or ""
    leito = header.get("leito") or ""
    data = header.get("data") or datetime.now().strftime("%d/%m/%Y")
    tipo = header.get("tipo") or "EVOLUÇÃO MÉDICA"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title=tipo,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["BodyText"], fontName="Helvetica",
                          fontSize=10, leading=13, alignment=TA_LEFT)
    head_style = ParagraphStyle("head", parent=styles["Normal"], fontSize=9,
                                textColor=HexColor("#555"))
    title_style = ParagraphStyle("title", parent=styles["Heading1"], fontSize=13,
                                 alignment=TA_CENTER, spaceAfter=8)
    meta_style = ParagraphStyle("meta", parent=styles["Normal"], fontSize=10,
                                alignment=TA_CENTER, spaceAfter=12)

    flow = []
    if hospital or medico:
        right = (medico or "") + (f" — CRM {crm}" if crm else "")
        flow.append(Paragraph(f"<b>{hospital}</b>" + (f"&nbsp;&nbsp;&nbsp;&nbsp;|&nbsp;&nbsp;&nbsp;&nbsp;{right}" if right else ""), head_style))
        flow.append(Spacer(1, 8))

    flow.append(Paragraph(tipo.upper(), title_style))
    meta_bits = []
    if paciente: meta_bits.append(f"<b>PACIENTE:</b> {paciente}")
    if leito: meta_bits.append(f"<b>LEITO:</b> {leito}")
    if data: meta_bits.append(f"<b>DATA:</b> {data}")
    if meta_bits:
        flow.append(Paragraph(" &middot; ".join(meta_bits), meta_style))
    flow.append(Paragraph("─" * 90, head_style))
    flow.append(Spacer(1, 8))

    for raw_line in text.split("\n"):
        safe = raw_line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if not safe.strip():
            flow.append(Spacer(1, 6))
        else:
            flow.append(Paragraph(safe, body))

    doc.build(flow)
    return buf.getvalue()


def _build_txt_bytes(text: str, header: Dict[str, Any]) -> bytes:
    hospital = header.get("hospital") or ""
    medico = header.get("medico") or ""
    crm = header.get("crm") or ""
    paciente = header.get("paciente") or ""
    leito = header.get("leito") or ""
    data = header.get("data") or datetime.now().strftime("%d/%m/%Y")
    tipo = header.get("tipo") or "EVOLUÇÃO MÉDICA"

    lines = []
    if hospital or medico:
        right = (medico or "") + (f" — CRM {crm}" if crm else "")
        lines.append(f"{hospital}    {right}".strip())
        lines.append("")
    lines.append(tipo.upper())
    meta = []
    if paciente: meta.append(f"PACIENTE: {paciente}")
    if leito: meta.append(f"LEITO: {leito}")
    if data: meta.append(f"DATA: {data}")
    if meta:
        lines.append("  ·  ".join(meta))
    lines.append("─" * 80)
    lines.append("")
    lines.append(text)
    return ("\n".join(lines)).encode("utf-8")


@app.post("/export-evolucao")
@app.post("/api/export/evolucao")
async def export_evolucao(req: ExportRequest):
    fmt = (req.format or "").lower().strip()
    if fmt not in {"docx", "pdf", "txt"}:
        raise HTTPException(status_code=400, detail="Formato deve ser docx, pdf ou txt.")
    if not req.text or not req.text.strip():
        raise HTTPException(status_code=400, detail="Texto vazio.")

    header = req.header or {}
    base = _safe_filename(req.filename_hint or header.get("paciente") or "evolucao")
    filename = f"{base}.{fmt}"

    try:
        if fmt == "docx":
            data = _build_docx_bytes(req.text, header)
            media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif fmt == "pdf":
            data = _build_pdf_bytes(req.text, header)
            media = "application/pdf"
        else:
            data = _build_txt_bytes(req.text, header)
            media = "text/plain; charset=utf-8"
    except Exception as e:
        print(f"Erro ao gerar {fmt}: {e}")
        raise HTTPException(status_code=500, detail=f"Falha ao gerar {fmt}: {e}")

    return StreamingResponse(
        io.BytesIO(data),
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# --- AI ENDPOINTS FOR EVOLUTION / PRESCRIPTION ---

class EvolutionRequest(BaseModel):
    patient: Dict[str, Any]
    tipo_unidade: str
    template: str
    data_plantao: str
    preferences: Optional[Dict[str, Any]] = None

@app.post("/api/ai/gerar-evolucao")
async def gerar_evolucao(req: EvolutionRequest):
    if not openai_key:
        raise HTTPException(status_code=500, detail="OpenAI API Key não configurada.")
    
    try:
        # Prompt build
        prompt = f"""
Você é um médico assistente sênior. Sua tarefa é gerar uma EVOLUÇÃO MÉDICA impecável baseada nos dados do paciente e no template fornecido.

DADOS DO PACIENTE:
{json.dumps(req.patient, indent=2, ensure_ascii=False)}

DATA DO PLANTÃO: {req.data_plantao}
TIPO DE UNIDADE: {req.tipo_unidade}

TEMPLATE OBRIGATÓRIO:
{req.template}

REGRAS:
1. Siga EXATAMENTE o template. Substitua os campos entre [colchetes] pelos dados reais.
2. Se não houver dado para um campo, use "SEM ALTERAÇÕES", "NADA A DECLARAR" ou remova o campo conforme o bom senso médico.
3. Use terminologia médica técnica e precisa.
4. Mantenha o texto em CAIXA ALTA se solicitado nas preferências.
5. Seja conciso e focado em dados relevantes.
"""
        
        pref = req.preferences or {}
        system_msg = "Você é um gerador de evoluções médicas técnicas. Retorne apenas o texto da evolução."
        if pref.get("uppercase"):
            system_msg += " SEMPRE EM CAIXA ALTA."

        response = await client.chat.completions.create(
            model=os.environ.get("OPENAI_MODEL", "gpt-4o-mini"),
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        
        text = response.choices[0].message.content
        return {"evolution_text": text}

    except Exception as e:
        print(f"Erro ao gerar evolução: {e}")
        raise HTTPException(status_code=500, detail=str(e))


class RefineEvolutionRequest(BaseModel):
    current_text: str
    instruction: str
    tipo_unidade: Optional[str] = None
    preferences: Optional[Dict[str, Any]] = None


@app.post("/api/ai/refinar-evolucao")
async def refinar_evolucao(req: RefineEvolutionRequest):
    if not openai_key:
        raise HTTPException(status_code=500, detail="OpenAI API Key não configurada.")
    if not req.current_text.strip():
        raise HTTPException(status_code=400, detail="Texto atual vazio.")
    if not req.instruction.strip():
        raise HTTPException(status_code=400, detail="Instrução vazia.")

    pref = req.preferences or {}
    uppercase = bool(pref.get("uppercase"))

    system_msg = (
        "Você edita evoluções médicas conforme instruções precisas. "
        "Retorne SOMENTE o texto final, sem comentários, sem markdown, sem cabeçalhos extra."
    )
    if uppercase:
        system_msg += " Use SEMPRE CAIXA ALTA."

    user_prompt = f"""Edite o texto abaixo conforme a instrução, preservando estrutura,
terminologia médica e qualquer informação que a instrução não mande remover.
Não invente dados clínicos novos — apenas reorganize, encurte, expanda ou corrija o
que está presente, segundo a instrução.

INSTRUÇÃO:
{req.instruction.strip()}

TEXTO ATUAL:
\"\"\"
{req.current_text}
\"\"\"

Retorne APENAS o texto editado."""

    try:
        response = await client.chat.completions.create(
            model=os.environ.get("OPENAI_MODEL", "gpt-4o-mini"),
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.0,
            max_tokens=16000,
        )
        text = response.choices[0].message.content or ""
        text = text.strip()
        if uppercase:
            text = text.upper()
        return {"evolution_text": text}
    except Exception as e:
        print(f"Erro ao refinar evolução: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/extract-async")
@app.post("/api/extract/extract-async")
async def extract_async(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    job_id = str(uuid.uuid4())
    file_bytes = await file.read()
    
    save_job(job_id, {
        "status": "queued",
        "stage": "Arquivo recebido",
        "file_name": file.filename,
        "result": None,
        "error": None
    })
    
    background_tasks.add_task(process_document_background, job_id, file_bytes, file.filename, file.content_type)
    
    return {
        "job_id": job_id,
        "status": "queued",
        "stage": "Arquivo recebido"
    }

@app.get("/job/{job_id}")
@app.get("/api/extract/job/{job_id}")
async def get_job_status(job_id: str):
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job não encontrado")
    return job

@app.delete("/cleanup-jobs")
@app.delete("/api/extract/cleanup-jobs")
async def cleanup_jobs():
    if supabase_client:
        try:
            # Delete jobs older than 24h
            threshold = (datetime.now() - timedelta(hours=24)).isoformat()
            supabase_client.table("extraction_jobs") \
                .delete() \
                .lt("created_at", threshold) \
                .execute()
            return {"message": "Jobs antigos removidos do Supabase"}
        except Exception as e:
            return {"message": f"Erro na limpeza do Supabase: {e}"}
    
    jobs_memory.clear()
    return {"message": "Memória local limpa"}

@app.get("/health")
@app.get("/api/health")
async def health_check():
    return {
        "status": "ok", 
        "service": "clinical-agents",
        "version": "async-extract-v3-e2e",
        "checks": {
            "supabase_service_key": bool(os.environ.get("SUPABASE_SERVICE_KEY")),
            "supabase_key": bool(os.environ.get("SUPABASE_KEY")),
            "supabase_url": bool(os.environ.get("SUPABASE_URL")),
            "openai_key": bool(openai_key),
            "openai_model_env": os.environ.get("OPENAI_MODEL")
        }
    }

if __name__ == "__main__":
    port_str = os.environ.get("PORT", "8000")
    try:
        # Tenta converter para int, removendo possíveis espaços ou caracteres extras
        if port_str.startswith('$'):
            print(f"Warning: Literal variable string detected in PORT: {port_str}")
            port = 8000
        else:
            port = int(port_str)
    except ValueError:
        print(f"Warning: Invalid PORT '{port_str}', defaulting to 8000")
        port = 8000
    
    print(f"Starting server on port {port}")
    uvicorn.run(app, host="0.0.0.0", port=port)
